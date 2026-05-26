import os
from pathlib import Path
from app.core.interfaces import BaseParser, ParsedDocument


class TxtParser(BaseParser):
    def parse(self, file_path: str) -> ParsedDocument:
        with open(file_path, encoding="utf-8") as f:
            text = f.read()
        return ParsedDocument(
            text=text,
            title=os.path.basename(file_path),
        )


class PdfParser(BaseParser):
    def parse(self, file_path: str) -> ParsedDocument:
        import fitz
        doc = fitz.open(file_path)
        pages = [page.get_text() for page in doc]
        text = "\n\n".join(pages)
        doc.close()
        return ParsedDocument(
            text=text,
            title=os.path.basename(file_path),
            metadata={"page_count": len(pages)},
            pages=pages,
        )


class WordParser(BaseParser):
    def parse(self, file_path: str) -> ParsedDocument:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n\n".join(paragraphs)
        return ParsedDocument(
            text=text,
            title=os.path.basename(file_path),
            metadata={"paragraph_count": len(paragraphs)},
        )


class MdParser(BaseParser):
    def parse(self, file_path: str) -> ParsedDocument:
        with open(file_path, encoding="utf-8") as f:
            text = f.read()
        return ParsedDocument(
            text=text,
            title=os.path.basename(file_path),
            metadata={"format": "markdown"},
        )


PARSER_REGISTRY: dict[str, type[BaseParser]] = {
    ".txt": TxtParser,
    ".pdf": PdfParser,
    ".docx": WordParser,
    ".doc": WordParser,
    ".md": MdParser,
    ".markdown": MdParser,
}


def get_parser(file_path: str) -> BaseParser:
    ext = Path(file_path).suffix.lower()
    parser_cls = PARSER_REGISTRY.get(ext)
    if parser_cls is None:
        raise ValueError(f"No parser registered for extension: {ext}")
    return parser_cls()
